"""Export Service - Generate PDF and Word formats"""
import logging
from typing import Optional
from io import BytesIO

logger = logging.getLogger(__name__)

class ExportService:
    """Generate PDF and Word document exports"""
    
    @staticmethod
    def markdown_to_pdf(markdown_content: str, filename: str = "test_plan.pdf") -> Optional[bytes]:
        """Convert Markdown to PDF"""
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
            from reportlab.lib.units import inch
            
            # Parse markdown headers and paragraphs (simple implementation)
            lines = markdown_content.split('\n')
            elements = []
            styles = getSampleStyleSheet()
            
            for line in lines:
                if line.startswith('# '):
                    title = line.replace('# ', '')
                    elements.append(Paragraph(title, styles['Heading1']))
                    elements.append(Spacer(1, 0.2*inch))
                elif line.startswith('## '):
                    heading = line.replace('## ', '')
                    elements.append(Paragraph(heading, styles['Heading2']))
                    elements.append(Spacer(1, 0.15*inch))
                elif line.startswith('### '):
                    subheading = line.replace('### ', '')
                    elements.append(Paragraph(subheading, styles['Heading3']))
                    elements.append(Spacer(1, 0.1*inch))
                elif line.startswith('- '):
                    bullet = line.replace('- ', '')
                    elements.append(Paragraph(f"• {bullet}", styles['Normal']))
                elif line.strip():
                    elements.append(Paragraph(line, styles['Normal']))
                else:
                    elements.append(Spacer(1, 0.1*inch))
            
            # Generate PDF
            pdf_buffer = BytesIO()
            doc = SimpleDocTemplate(pdf_buffer, pagesize=letter)
            doc.build(elements)
            
            return pdf_buffer.getvalue()
        except ImportError:
            logger.warning("reportlab not installed")
            return None
        except Exception as e:
            logger.error(f"PDF generation error: {e}")
            return None
    
    @staticmethod
    def markdown_to_docx(markdown_content: str, filename: str = "test_plan.docx") -> Optional[bytes]:
        """Convert Markdown to Word .docx"""
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor, Inches
            
            doc = Document()
            lines = markdown_content.split('\n')
            
            for line in lines:
                if line.startswith('# '):
                    title = line.replace('# ', '')
                    heading = doc.add_heading(title, level=1)
                elif line.startswith('## '):
                    heading = line.replace('## ', '')
                    doc.add_heading(heading, level=2)
                elif line.startswith('### '):
                    subheading = line.replace('### ', '')
                    doc.add_heading(subheading, level=3)
                elif line.startswith('- '):
                    bullet = line.replace('- ', '')
                    doc.add_paragraph(bullet, style='List Bullet')
                elif line.startswith('* '):
                    bullet = line.replace('* ', '')
                    doc.add_paragraph(bullet, style='List Bullet')
                elif line.strip():
                    doc.add_paragraph(line)
                else:
                    # Empty line
                    if doc.paragraphs:
                        doc.add_paragraph()
            
            # Save to bytes
            docx_buffer = BytesIO()
            doc.save(docx_buffer)
            docx_buffer.seek(0)
            
            return docx_buffer.getvalue()
        except ImportError:
            logger.warning("python-docx not installed")
            return None
        except Exception as e:
            logger.error(f"DOCX generation error: {e}")
            return None
